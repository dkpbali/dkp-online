"""
generate_ba.py  –  Engine generate Berita Acara dari template .docx
"""
import base64, io
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Cm

TEMPLATE = Path(__file__).parent / "templates" / "Berita_Acara_Template.docx"

HARI_ID  = ["Senin","Selasa","Rabu","Kamis","Jumat","Sabtu","Minggu"]
BULAN_ID = ["","Januari","Februari","Maret","April","Mei","Juni",
            "Juli","Agustus","September","Oktober","November","Desember"]


def _fmt_tanggal(val) -> str:
    if isinstance(val, datetime):
        obj = val
    elif isinstance(val, str):
        for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y"):
            try: obj = datetime.strptime(val, fmt); break
            except: obj = None
        if not obj: return str(val)
    else:
        return str(val) if val else ""
    return f"{obj.day} {BULAN_ID[obj.month]} {obj.year}"


def _hari(val) -> str:
    if isinstance(val, datetime): return HARI_ID[val.weekday()]
    for fmt in ("%d/%m/%Y", "%Y-%m-%d"):
        try: return HARI_ID[datetime.strptime(val, fmt).weekday()]
        except: pass
    return ""


def _replace_text(doc, key, value):
    tag = "{{" + key + "}}"
    def _in_para(para):
        full = "".join(r.text for r in para.runs)
        if tag not in full: return
        replaced = full.replace(tag, str(value))
        for i, run in enumerate(para.runs):
            run.text = replaced if i == 0 else ""
    for p in doc.paragraphs: _in_para(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs: _in_para(p)


def _replace_ttd(doc, img_bytes_or_path):
    tag = "{{TTD}}"
    def _in_para(para):
        full = "".join(r.text for r in para.runs)
        if tag not in full: return False
        for run in para.runs: run.text = ""
        src = io.BytesIO(img_bytes_or_path) if isinstance(img_bytes_or_path, bytes) \
              else str(img_bytes_or_path)
        para.runs[0].add_picture(src, width=Cm(4))
        return True
    for p in doc.paragraphs:
        if _in_para(p): return
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _in_para(p): return


def generate(data: dict, ttd=None, template_path=None) -> bytes:
    """
    data  : dict dengan key = nama placeholder (tanpa {{ }})
    ttd   : bytes PNG/JPG  ATAU  str/Path ke file gambar
    return: bytes file .docx
    """
    tpl = Path(template_path or TEMPLATE)
    doc = Document(str(tpl))

    # Auto-derive HARI & format TANGGAL_PINJAM
    tgl = data.get("TANGGAL_PINJAM", "")
    data.setdefault("HARI", _hari(tgl))
    data["TANGGAL_PINJAM"] = _fmt_tanggal(tgl)

    for key, val in data.items():
        if key != "TTD":
            _replace_text(doc, key, val or "")

    if ttd:
        _replace_ttd(doc, ttd)
    else:
        _replace_text(doc, "TTD", "_____________________")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
